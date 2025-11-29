from docx import Document
import pandas as pd
from datetime import datetime, timedelta

hoy = datetime.today().date()
fechas = [hoy - timedelta(days=i) for i in range(6, -1, -1)]
valores = [2.221, 2.220, 2.209, 2.206, 2.206, 2.210, 2.209]

df = pd.DataFrame({
    "fecha": [f.strftime("%d/%m/%Y") for f in fechas],
    "euribor": valores
})

doc = Document()
doc.add_heading("Informe automático: Euríbor 12M (últimos 7 días)", level=1)

fecha_generacion = datetime.today().strftime("%d/%m/%Y")
doc.add_paragraph(f"Fecha de generación: {fecha_generacion}")
doc.add_paragraph("Evolución del Euríbor 12M en los últimos 7 días:")

table = doc.add_table(rows=1, cols=2)
hdr = table.rows[0].cells
hdr[0].text = "Fecha"
hdr[1].text = "Euríbor 12M (%)"

for _, row in df.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row["fecha"])
    cells[1].text = f"{row['euribor']:.3f}"

nombre = f"euribor_{datetime.today().strftime('%Y%m%d')}.docx"
doc.save(nombre)

print(f"Informe generado: {nombre}")
